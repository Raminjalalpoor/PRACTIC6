from __future__ import annotations

import html
import shutil
import subprocess
import sys
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT))

from data_processing import analyse, read_sales_file


DOCS = ROOT / "docs"
SCREENSHOTS = DOCS / "screenshots"
HTML_REPORT = DOCS / "practice_report.html"
DOCX_REPORT = DOCS / "practice_report.docx"
PDF_REPORT = DOCS / "practice_report.pdf"


def rubles(value: float) -> str:
    return f"{value:,.0f}".replace(",", " ") + " ₽"


def number(value: float) -> str:
    return f"{value:,.0f}".replace(",", " ")


def today_ru() -> str:
    return date.today().strftime("%d.%m.%Y")


def chrome_path() -> str | None:
    candidates = [
        shutil.which("chrome"),
        shutil.which("msedge"),
        r"C:\Program Files\Google\Chrome\Application\chrome.exe",
        r"C:\Program Files (x86)\Google\Chrome\Application\chrome.exe",
        r"C:\Program Files\Microsoft\Edge\Application\msedge.exe",
        r"C:\Program Files (x86)\Microsoft\Edge\Application\msedge.exe",
    ]
    for item in candidates:
        if item and Path(item).exists():
            return item
    return None


def collect_report_data() -> dict:
    frame = read_sales_file(ROOT / "data" / "sample_sales.csv")
    analysis = analyse(frame)
    metrics = analysis.metrics
    peak = analysis.seasonality.get("peak_month") or "не определен"

    return {
        "metrics": [
            ("Общая выручка", rubles(metrics["revenue"])),
            ("Продано единиц", number(metrics["units"])),
            ("Количество записей", number(metrics["records"])),
            ("Уникальных товаров", number(metrics["unique_products"])),
            ("Средняя операция", rubles(metrics["average_operation"])),
            ("Популярный товар", metrics["popular_product"]),
            ("Товар-лидер по выручке", metrics["revenue_leader"]),
            ("Пиковый месяц спроса", peak.capitalize()),
        ],
        "top_products": analysis.products.head(5).to_dict("records"),
        "monthly": analysis.monthly[["Месяц", "Количество", "Выручка"]].tail(6).to_dict("records"),
        "forecast": analysis.forecast,
        "seasonality": analysis.seasonality,
    }


def add_docx_table(document: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    for index, header in enumerate(headers):
        header_cells[index].text = header
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = str(value)


def add_docx_heading(document: Document, text: str, level: int = 1) -> None:
    paragraph = document.add_heading(text, level=level)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT


def build_docx(data: dict) -> None:
    document = Document()
    section = document.sections[0]
    section.top_margin = Cm(1.7)
    section.bottom_margin = Cm(1.7)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(1.5)

    styles = document.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(14)

    def cover_line(
        text: str = "",
        *,
        bold: bool = False,
        size: int = 14,
        align: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.CENTER,
        before: int = 0,
        after: int = 0,
    ):
        paragraph = document.add_paragraph()
        paragraph.alignment = align
        paragraph.paragraph_format.space_before = Pt(before)
        paragraph.paragraph_format.space_after = Pt(after)
        run = paragraph.add_run(text)
        run.bold = bold
        run.font.name = "Times New Roman"
        run.font.size = Pt(size)
        return paragraph

    cover_line("Министерство науки и высшего образования Российской Федерации", bold=True, after=18)
    cover_line("ФЕДЕРАЛЬНОЕ ГОСУДАРСТВЕННОЕ АВТОНОМНОЕ", bold=True)
    cover_line("ОБРАЗОВАТЕЛЬНОЕ УЧРЕЖДЕНИЕ ВЫСШЕГО ОБРАЗОВАНИЯ", bold=True)
    cover_line("НАЦИОНАЛЬНЫЙ ИССЛЕДОВАТЕЛЬСКИЙ УНИВЕРСИТЕТ ИТМО", bold=True)
    cover_line("(Университет ИТМО)", after=22)
    cover_line("Факультет систем управления и робототехники")

    cover_line(before=76)
    cover_line("ОТЧЕТ ПО ПРАКТИКЕ", bold=True, size=16, after=18)
    cover_line("по теме:", after=8)
    cover_line(
        "РАЗРАБОТКА ВЕБ-ПРИЛОЖЕНИЯ ДЛЯ АНАЛИЗА ПРОДАЖ МАГАЗИНА,",
        bold=True,
        size=14,
    )
    cover_line(
        "ВИЗУАЛИЗАЦИИ РЕЗУЛЬТАТОВ И ПРОГНОЗИРОВАНИЯ СПРОСА",
        bold=True,
        size=14,
    )
    cover_line("Shop Analytics RU", bold=True, size=14, after=52)

    cover_line("Выполнил:", align=WD_ALIGN_PARAGRAPH.RIGHT)
    cover_line("Махенге Антон Бинилит 340293", align=WD_ALIGN_PARAGRAPH.RIGHT)
    cover_line("Группа: R3337", align=WD_ALIGN_PARAGRAPH.RIGHT, after=20)
    cover_line("Преподаватель:", align=WD_ALIGN_PARAGRAPH.RIGHT)
    cover_line("к.т.н., преподаватель ФСУиР", align=WD_ALIGN_PARAGRAPH.RIGHT)
    cover_line("Рассадина Анна Александровна", align=WD_ALIGN_PARAGRAPH.RIGHT, after=48)

    cover_line("Санкт-Петербург")
    cover_line("2026")

    document.add_page_break()

    add_docx_heading(document, "1. Цель проекта")
    document.add_paragraph(
        "Целью работы стала разработка русскоязычного веб-приложения для загрузки CSV/XLSX-файлов "
        "с продажами магазина, автоматического расчета ключевых показателей, построения интерактивных "
        "графиков, фильтрации данных, прогноза спроса и экспорта аналитического отчета в Excel."
    )

    add_docx_heading(document, "2. Основные задачи")
    for item in [
        "Создать Flask-приложение с загрузкой файлов и современным интерфейсом.",
        "Поддержать CSV из Excel: разделители запятая/точка с запятой, UTF-8 и Windows-1251, десятичную запятую.",
        "Проверять обязательные столбцы: Дата, Товар, Количество, Цена; поддержать опциональную Категорию.",
        "Рассчитывать выручку, количество продаж, ассортимент, среднюю операцию, популярный товар и лидера по выручке.",
        "Построить графики Plotly: выручка по месяцам, ТОП-10 товаров и доля выручки по категориям.",
        "Добавить фильтрацию по дате и товару с пересчетом всех KPI и графиков.",
        "Реализовать прогноз спроса на следующий месяц через линейную регрессию NumPy.",
        "Сформировать Excel-отчет с листами Показатели, Продажи, По товарам, По месяцам и Прогноз.",
    ]:
        document.add_paragraph(item, style="List Bullet")

    add_docx_heading(document, "3. Структура проекта")
    add_docx_table(
        document,
        ["Файл или папка", "Назначение"],
        [
            ["app.py", "Создание Flask-приложения, настройки, фильтры форматирования."],
            ["routes.py", "Маршруты загрузки, дашборда, экспорта и демонстрационного файла."],
            ["data_processing.py", "Чтение, валидация, агрегация, прогноз и сезонность."],
            ["visualization.py", "Формирование интерактивных Plotly-графиков."],
            ["reporting.py", "Экспорт Excel-отчета."],
            ["templates/", "HTML-шаблоны главной страницы и дашборда."],
            ["static/", "CSS и JavaScript интерфейса."],
            ["data/sample_sales.csv", "Демонстрационный набор данных российского магазина."],
            ["tests/", "Автоматические проверки загрузки, аналитики, ошибок и экспорта."],
        ],
    )

    add_docx_heading(document, "4. Реализованный пользовательский сценарий")
    for item in [
        "Пользователь открывает главную страницу приложения.",
        "Выбирает CSV или XLSX-файл с продажами и нажимает кнопку анализа.",
        "Система сохраняет файл во временную папку uploads под UUID и привязывает его к сессии.",
        "После проверки данных открывается дашборд с KPI-карточками, таблицами и графиками.",
        "Пользователь применяет фильтр по периоду и товару; показатели и визуализации пересчитываются.",
        "По кнопке Скачать отчет Excel формируется отчет с учетом текущих фильтров.",
    ]:
        document.add_paragraph(item, style="List Number")

    add_docx_heading(document, "5. Контрольные показатели на демонстрационных данных")
    add_docx_table(document, ["Показатель", "Значение"], [[name, value] for name, value in data["metrics"]])

    add_docx_heading(document, "6. ТОП товаров")
    add_docx_table(
        document,
        ["Товар", "Количество", "Выручка"],
        [
            [row["Товар"], number(row["Количество"]), rubles(row["Выручка"])]
            for row in data["top_products"]
        ],
    )

    add_docx_heading(document, "7. Прогноз и сезонность")
    forecast = data["forecast"]
    if forecast["available"]:
        document.add_paragraph(
            f"Прогноз на {forecast['next_month']}: около {number(forecast['total'])} единиц продаж. "
            "Расчет выполнен методом линейной регрессии NumPy по месячному временному ряду."
        )
    else:
        document.add_paragraph(f"Прогноз недоступен: {forecast['message']}")
    for recommendation in data["seasonality"].get("recommendations", []):
        document.add_paragraph(recommendation, style="List Bullet")

    add_docx_heading(document, "8. Проверка качества")
    for item in [
        "Создано виртуальное окружение Python 3.14.",
        "Запущены автоматические тесты: 11 тестов, результат OK.",
        "Проверены CSV UTF-8, Windows-1251, XLSX, пропущенные столбцы, ошибочные даты, текст вместо чисел и отрицательные значения.",
        "Проверены загрузка файла, фильтрация, построение 3 графиков Plotly и скачивание Excel в браузере Chrome.",
        "Созданы скриншоты главной страницы, дашборда и отфильтрованного дашборда.",
    ]:
        document.add_paragraph(item, style="List Bullet")

    add_docx_heading(document, "9. Скриншоты приложения")
    screenshots = [
        ("Главная страница загрузки файла", SCREENSHOTS / "01_home.png"),
        ("Дашборд после загрузки данных", SCREENSHOTS / "02_dashboard.png"),
        ("Дашборд после фильтрации по периоду и товару", SCREENSHOTS / "03_filtered_dashboard.png"),
    ]
    for caption, path in screenshots:
        if path.exists():
            document.add_paragraph(caption).runs[0].bold = True
            document.add_picture(str(path), width=Cm(16.5))

    add_docx_heading(document, "10. Вывод")
    document.add_paragraph(
        "В ходе работы разработано законченное учебное веб-приложение Shop Analytics RU. "
        "Приложение соответствует техническому заданию: принимает файлы продаж, проверяет данные, "
        "строит аналитический дашборд, поддерживает российские форматы чисел и дат, прогнозирует спрос "
        "и экспортирует отчет. Проект можно запускать локально без базы данных и использовать как основу "
        "для дальнейшего развития: добавления авторизации, базы данных, чеков, возвратов и расширенных моделей прогнозирования."
    )

    document.add_section(WD_SECTION.NEW_PAGE)
    add_docx_heading(document, "Приложение. Команды запуска")
    add_docx_table(
        document,
        ["Действие", "Команда"],
        [
            ["Создать окружение", "py -3.14 -m venv .venv"],
            ["Активировать окружение", r".\.venv\Scripts\activate"],
            ["Установить зависимости", "pip install -r requirements.txt"],
            ["Запустить приложение", "python app.py"],
            ["Запустить тесты", "python -m unittest discover -s tests -v"],
        ],
    )

    document.save(DOCX_REPORT)


def build_html(data: dict) -> None:
    def table(headers: list[str], rows: list[list[str]]) -> str:
        head = "".join(f"<th>{html.escape(header)}</th>" for header in headers)
        body = "".join(
            "<tr>" + "".join(f"<td>{html.escape(str(value))}</td>" for value in row) + "</tr>"
            for row in rows
        )
        return f"<table><thead><tr>{head}</tr></thead><tbody>{body}</tbody></table>"

    screenshots = [
        ("Главная страница загрузки файла", SCREENSHOTS / "01_home.png"),
        ("Дашборд после загрузки данных", SCREENSHOTS / "02_dashboard.png"),
        ("Дашборд после фильтрации по периоду и товару", SCREENSHOTS / "03_filtered_dashboard.png"),
    ]
    screenshot_html = "\n".join(
        f"<figure><img src='{path.as_uri()}' alt='{html.escape(caption)}'><figcaption>{html.escape(caption)}</figcaption></figure>"
        for caption, path in screenshots
        if path.exists()
    )

    recommendations = "".join(
        f"<li>{html.escape(item)}</li>" for item in data["seasonality"].get("recommendations", [])
    )
    forecast = data["forecast"]
    forecast_text = (
        f"Прогноз на {html.escape(forecast['next_month'])}: около {number(forecast['total'])} единиц продаж."
        if forecast["available"]
        else f"Прогноз недоступен: {html.escape(forecast['message'])}"
    )

    content = f"""<!doctype html>
<html lang="ru">
<head>
  <meta charset="utf-8">
  <title>Отчет по практике — Shop Analytics RU</title>
  <style>
    @page {{ size: A4; margin: 16mm 14mm; }}
    body {{ font-family: "Times New Roman", serif; color: #111; line-height: 1.35; }}
    h2 {{ margin-top: 28px; padding-bottom: 6px; border-bottom: 2px solid #10233f; font-size: 20px; }}
    h3 {{ margin-top: 20px; font-size: 16px; }}
    p, li, td, th {{ font-size: 12px; }}
    .cover {{ min-height: 265mm; position: relative; page-break-after: always; text-align: center; }}
    .cover-top {{ padding-top: 20mm; }}
    .cover-top p {{ margin: 0 0 7px; font-size: 15px; }}
    .cover-top .bold {{ font-weight: 700; }}
    .cover-work {{ margin-top: 48mm; }}
    .cover-work p {{ margin: 0 0 12px; font-size: 16px; }}
    .cover-work .title {{ font-size: 18px; font-weight: 700; }}
    .cover-work .topic {{ margin-bottom: 2px; font-size: 16px; font-weight: 700; text-transform: uppercase; }}
    .cover-side {{ width: 92mm; margin: 34mm 4mm 0 auto; text-align: right; }}
    .cover-side p {{ margin: 0 0 7px; font-size: 15px; }}
    .cover-bottom {{ position: absolute; bottom: 9mm; left: 0; right: 0; }}
    .cover-bottom p {{ margin: 0 0 5px; font-size: 15px; }}
    .page-break {{ page-break-before: always; }}
    table {{ width: 100%; border-collapse: collapse; margin: 12px 0 18px; }}
    th {{ color: white; background: #10233f; }}
    th, td {{ border: 1px solid #d6dee9; padding: 8px 9px; text-align: left; vertical-align: top; }}
    tr:nth-child(even) td {{ background: #f7f9fc; }}
    .callout {{ padding: 13px 15px; background: #eaf8f3; border-left: 4px solid #10a979; border-radius: 8px; }}
    figure {{ page-break-inside: avoid; margin: 18px 0 26px; }}
    figcaption {{ margin-top: 8px; color: #506071; font-size: 11px; font-weight: 700; }}
    img {{ width: 100%; border: 1px solid #d6dee9; border-radius: 8px; }}
    code {{ color: #0f5132; background: #eef8f3; padding: 2px 4px; border-radius: 4px; }}
  </style>
</head>
<body>
  <div class="cover">
    <div class="cover-top">
      <p class="bold">Министерство науки и высшего образования Российской Федерации</p>
      <p class="bold">ФЕДЕРАЛЬНОЕ ГОСУДАРСТВЕННОЕ АВТОНОМНОЕ</p>
      <p class="bold">ОБРАЗОВАТЕЛЬНОЕ УЧРЕЖДЕНИЕ ВЫСШЕГО ОБРАЗОВАНИЯ</p>
      <p class="bold">НАЦИОНАЛЬНЫЙ ИССЛЕДОВАТЕЛЬСКИЙ УНИВЕРСИТЕТ ИТМО</p>
      <p>(Университет ИТМО)</p>
      <p style="margin-top: 18px;">Факультет систем управления и робототехники</p>
    </div>

    <div class="cover-work">
      <p class="title">ОТЧЕТ ПО ПРАКТИКЕ</p>
      <p>по теме:</p>
      <p class="topic">РАЗРАБОТКА ВЕБ-ПРИЛОЖЕНИЯ ДЛЯ АНАЛИЗА ПРОДАЖ МАГАЗИНА,</p>
      <p class="topic">ВИЗУАЛИЗАЦИИ РЕЗУЛЬТАТОВ И ПРОГНОЗИРОВАНИЯ СПРОСА</p>
      <p class="topic">Shop Analytics RU</p>
    </div>

    <div class="cover-side">
      <p>Выполнил:</p>
      <p>Махенге Антон Бинилит 340293</p>
      <p>Группа: R3337</p>
      <p style="margin-top: 20px;">Преподаватель:</p>
      <p>к.т.н., преподаватель ФСУиР</p>
      <p>Рассадина Анна Александровна</p>
    </div>

    <div class="cover-bottom">
      <p>Санкт-Петербург</p>
      <p>2026</p>
    </div>
  </div>

  <h2>1. Цель проекта</h2>
  <p>Целью работы стала разработка русскоязычного веб-приложения для загрузки CSV/XLSX-файлов с продажами магазина, автоматического расчета ключевых показателей, построения интерактивных графиков, фильтрации данных, прогноза спроса и экспорта аналитического отчета в Excel.</p>

  <h2>2. Основные задачи</h2>
  <ul>
    <li>Создать Flask-приложение с загрузкой файлов и современным интерфейсом.</li>
    <li>Поддержать CSV из Excel: разделители запятая/точка с запятой, UTF-8 и Windows-1251, десятичную запятую.</li>
    <li>Проверять обязательные столбцы: Дата, Товар, Количество, Цена; поддержать опциональную Категорию.</li>
    <li>Рассчитывать KPI, строить графики Plotly, выполнять фильтрацию и экспорт отчета.</li>
    <li>Реализовать прогноз спроса через линейную регрессию NumPy и осторожные сезонные рекомендации.</li>
  </ul>

  <h2>3. Структура проекта</h2>
  {table(["Файл или папка", "Назначение"], [
      ["app.py", "Создание Flask-приложения, настройки, фильтры форматирования."],
      ["routes.py", "Маршруты загрузки, дашборда, экспорта и демонстрационного файла."],
      ["data_processing.py", "Чтение, валидация, агрегация, прогноз и сезонность."],
      ["visualization.py", "Формирование интерактивных Plotly-графиков."],
      ["reporting.py", "Экспорт Excel-отчета."],
      ["templates/ и static/", "Шаблоны, стили и JavaScript интерфейса."],
      ["tests/", "Автоматические проверки приложения."],
  ])}

  <h2>4. Реализованный сценарий работы</h2>
  <ol>
    <li>Пользователь открывает главную страницу приложения.</li>
    <li>Загружает CSV или XLSX-файл с продажами.</li>
    <li>Система сохраняет файл во временную папку uploads под UUID и привязывает его к Flask-сессии.</li>
    <li>После проверки данных открывается дашборд с KPI-карточками, таблицами и графиками.</li>
    <li>Фильтр по датам и товару пересчитывает показатели и визуализации.</li>
    <li>Кнопка скачивания формирует Excel-отчет по текущему дашборду.</li>
  </ol>

  <h2>5. Контрольные показатели</h2>
  {table(["Показатель", "Значение"], [[name, value] for name, value in data["metrics"]])}

  <h2>6. ТОП товаров</h2>
  {table(["Товар", "Количество", "Выручка"], [[row["Товар"], number(row["Количество"]), rubles(row["Выручка"])] for row in data["top_products"]])}

  <h2>7. Прогноз и сезонность</h2>
  <p class="callout">{forecast_text}</p>
  <ul>{recommendations}</ul>

  <h2>8. Проверка качества</h2>
  <ul>
    <li>Виртуальное окружение создано через <code>py -3.14 -m venv .venv</code>.</li>
    <li>Автоматические тесты: 11 тестов, результат OK.</li>
    <li>Проверены CSV UTF-8, Windows-1251, XLSX, ошибки столбцов, дат и числовых значений.</li>
    <li>Через Chrome проверены загрузка файла, дашборд, фильтр, 3 Plotly-графика и скачивание Excel.</li>
  </ul>

  <h2>9. Скриншоты приложения</h2>
  {screenshot_html}

  <h2>10. Вывод</h2>
  <p>В ходе работы разработано законченное учебное веб-приложение Shop Analytics RU. Приложение соответствует техническому заданию: принимает файлы продаж, проверяет данные, строит аналитический дашборд, поддерживает российские форматы чисел и дат, прогнозирует спрос и экспортирует отчет. Проект можно запускать локально без базы данных и использовать как основу для дальнейшего развития.</p>

  <h2>Приложение. Команды запуска</h2>
  {table(["Действие", "Команда"], [
      ["Создать окружение", "py -3.14 -m venv .venv"],
      ["Активировать окружение", r".\\.venv\\Scripts\\activate"],
      ["Установить зависимости", "pip install -r requirements.txt"],
      ["Запустить приложение", "python app.py"],
      ["Запустить тесты", "python -m unittest discover -s tests -v"],
  ])}
</body>
</html>"""
    HTML_REPORT.write_text(content, encoding="utf-8")


def build_pdf() -> None:
    chrome = chrome_path()
    if chrome is None:
        raise RuntimeError("Chrome или Edge не найден. PDF можно создать вручную из HTML.")
    subprocess.run(
        [
            chrome,
            "--headless",
            "--disable-gpu",
            "--no-sandbox",
            "--no-pdf-header-footer",
            f"--print-to-pdf={PDF_REPORT}",
            HTML_REPORT.as_uri(),
        ],
        cwd=ROOT,
        check=True,
        capture_output=True,
        text=True,
        encoding="utf-8",
        errors="replace",
    )


def main() -> None:
    DOCS.mkdir(exist_ok=True)
    data = collect_report_data()
    build_docx(data)
    build_html(data)
    build_pdf()
    print(f"Создан Word-отчет: {DOCX_REPORT}")
    print(f"Создан PDF-отчет:  {PDF_REPORT}")
    print(f"Создан HTML-исходник: {HTML_REPORT}")


if __name__ == "__main__":
    main()
